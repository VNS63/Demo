import io
from typing import Optional

from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from fastapi import FastAPI, HTTPException, Request
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, Field

app = FastAPI()


class HtmlToDocxRequest(BaseModel):
    html_content: str = Field(..., min_length=1, description="HTML string to convert")
    filename: Optional[str] = Field(
        default="converted.docx", description="Name of the generated DOCX file"
    )


def _safe_filename(filename: Optional[str]) -> str:
    if not filename:
        return "converted.docx"

    cleaned = filename.strip().replace("/", "_").replace("\\", "_")
    if not cleaned.lower().endswith(".docx"):
        cleaned += ".docx"

    return cleaned or "converted.docx"


def _add_inline_content(paragraph, node) -> None:
    if isinstance(node, NavigableString):
        text = str(node)
        if text:
            paragraph.add_run(text)
        return

    if not isinstance(node, Tag):
        return

    if node.name == "br":
        paragraph.add_run("\n")
        return

    if node.name in {"strong", "b"}:
        run = paragraph.add_run(node.get_text())
        run.bold = True
        return

    if node.name in {"em", "i"}:
        run = paragraph.add_run(node.get_text())
        run.italic = True
        return

    if node.name == "u":
        run = paragraph.add_run(node.get_text())
        run.underline = True
        return

    for child in node.children:
        _add_inline_content(paragraph, child)


def _add_block_content(document: Document, element: Tag) -> None:
    name = element.name.lower()

    if name in {"h1", "h2", "h3", "h4", "h5", "h6"}:
        level = int(name[1])
        heading = document.add_heading(level=level)
        for child in element.children:
            _add_inline_content(heading, child)
        return

    if name in {"p", "div", "article", "section"}:
        paragraph = document.add_paragraph()
        for child in element.children:
            _add_inline_content(paragraph, child)
        return

    if name in {"ul", "ol"}:
        style = "List Bullet" if name == "ul" else "List Number"
        for li in element.find_all("li", recursive=False):
            paragraph = document.add_paragraph(style=style)
            for child in li.children:
                _add_inline_content(paragraph, child)
        return

    if name == "pre":
        document.add_paragraph(element.get_text())
        return

    if name == "table":
        rows = element.find_all("tr")
        if not rows:
            return
        first_row_cells = rows[0].find_all(["th", "td"])
        table = document.add_table(rows=len(rows), cols=max(1, len(first_row_cells)))
        for row_idx, row in enumerate(rows):
            cells = row.find_all(["th", "td"])
            for col_idx, cell in enumerate(cells):
                if col_idx < len(table.rows[row_idx].cells):
                    table.rows[row_idx].cells[col_idx].text = cell.get_text(" ", strip=True)
        return

    if name in {"script", "style"}:
        return

    # Fallback for unknown containers: recurse into children.
    for child in element.children:
        if isinstance(child, Tag):
            _add_block_content(document, child)
        elif isinstance(child, NavigableString) and str(child).strip():
            document.add_paragraph(str(child).strip())


def _html_to_docx_bytes(html_content: str) -> io.BytesIO:
    soup = BeautifulSoup(html_content, "html.parser")
    document = Document()
    root = soup.body or soup

    for node in root.children:
        if isinstance(node, Tag):
            _add_block_content(document, node)
        elif isinstance(node, NavigableString) and str(node).strip():
            document.add_paragraph(str(node).strip())

    output = io.BytesIO()
    document.save(output)
    output.seek(0)
    return output


async def _extract_html_payload(request: Request) -> tuple[str, Optional[str]]:
    content_type = (request.headers.get("content-type") or "").lower()

    if "application/json" in content_type:
        data = await request.json()
        if isinstance(data, dict):
            return str(data.get("html_content", "")), data.get("filename")
        if isinstance(data, str):
            return data, None
        return "", None

    if "multipart/form-data" in content_type or "application/x-www-form-urlencoded" in content_type:
        form = await request.form()
        html_content = form.get("html_content")
        filename = form.get("filename")
        return str(html_content or ""), str(filename) if filename else None

    body = await request.body()
    return body.decode("utf-8", errors="replace"), None


@app.post("/convert-html-to-docx")
async def convert_html_to_docx(payload: HtmlToDocxRequest):
    try:
        file_stream = _html_to_docx_bytes(payload.html_content)
        output_name = _safe_filename(payload.filename)

        return StreamingResponse(
            file_stream,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{output_name}"'},
        )
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Conversion failed: {exc}") from exc


@app.post("/convert-html-file-to-docx")
async def convert_html_file_to_docx(request: Request):
    try:
        html_content, filename = await _extract_html_payload(request)
        if not html_content.strip():
            raise HTTPException(
                status_code=400,
                detail=(
                    "No HTML content provided. Send JSON {'html_content': '...'} "
                    "or form field html_content, or raw text/html body."
                ),
            )

        file_stream = _html_to_docx_bytes(html_content)
        output_name = _safe_filename(filename)

        return StreamingResponse(
            file_stream,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{output_name}"'},
        )
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Conversion failed: {exc}") from exc
